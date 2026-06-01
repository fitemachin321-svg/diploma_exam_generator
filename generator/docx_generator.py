from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

class ExamDOCXGenerator:
    def __init__(self, filename, title):
        self.filename = filename
        self.title = title
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        try:
            normal_style = self.doc.styles['Normal']
            normal_font = normal_style.font
            normal_font.name = 'Times New Roman'
            normal_font.size = Pt(14)
        except:
            pass
        
        try:
            header_style = self.doc.styles['Heading 1']
            header_font = header_style.font
            header_font.name = 'Times New Roman'
            header_font.size = Pt(16)
            header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass
        
        try:
            variant_style = self.doc.styles['Heading 2']
            variant_font = variant_style.font
            variant_font.name = 'Times New Roman'
            variant_font.size = Pt(14)
        except:
            pass
    
    def _add_header_info(self):
        date_str = datetime.now().strftime("%d.%m.%Y %H:%M")
        info_paragraph = self.doc.add_paragraph()
        info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_paragraph.add_run(f"Дата генерации: {date_str}")
        info_run.font.size = Pt(11)
        info_run.font.name = 'Times New Roman'
    
    def generate_header(self):
        header = self.doc.add_heading(self.title, level=1)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_header_info()
        self.doc.add_paragraph()
    
    def add_variant(self, variant_number, questions):
        self.doc.add_heading(f"Вариант {variant_number}", level=2)
        for i, q in enumerate(questions, 1):
            q_text = f"{i}. {q.get('text', 'Нет текста')}"
            if q.get('points'):
                q_text += f" (макс. {q.get('points')} баллов)"
            p = self.doc.add_paragraph(q_text)
            p.paragraph_format.space_after = Pt(12)
            answer_p = self.doc.add_paragraph("Ответ: ____________________")
            answer_p.paragraph_format.space_after = Pt(24)
        self.doc.add_paragraph()
    
    def generate_answers(self, answers):
        self.doc.add_heading("ОТВЕТЫ (для учителя)", level=2)
        for variant_num in sorted(answers.keys()):
            self.doc.add_heading(f"Вариант {variant_num}:", level=3)
            for i, ans in enumerate(answers[variant_num], 1):
                p = self.doc.add_paragraph(f"{i}. {ans}")
                p.paragraph_format.space_after = Pt(6)
            self.doc.add_paragraph()
    
    def build(self):
        self.doc.save(self.filename)